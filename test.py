from docx import Document

doc = Document()
doc.add_heading("Arduino Quorum Sensing Circuit - Wiring and Pin Details", level=1)

doc.add_paragraph(
    "This document describes all wiring connections and pin mappings for the Arduino-based Quorum Sensing Simulation project. "
    "It includes every node, LED, servo motor, buzzer, LDR, and reset button."
)

# Component list
doc.add_heading("Component List", level=2)
components = [
    "Arduino Uno board",
    "3 × Potentiometers (10 kΩ)",
    "3 × LEDs with 220 Ω resistors",
    "3 × Servo motors (SG90 type)",
    "1 × LDR + 10 kΩ resistor",
    "1 × Push button (reset)",
    "1 × Buzzer",
    "Breadboard + jumper wires"
]
for c in components:
    doc.add_paragraph(c, style="List Bullet")

# Wiring table
doc.add_heading("Pin Connections", level=2)
table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
hdr[0].text = "Component"
hdr[1].text = "Pin/Leg"
hdr[2].text = "Arduino Pin"
hdr[3].text = "Description"

connections = [
    ("Potentiometer 1", "Middle leg", "A0", "AI input for Node 1"),
    ("Potentiometer 2", "Middle leg", "A1", "AI input for Node 2"),
    ("Potentiometer 3", "Middle leg", "A2", "AI input for Node 3"),
    ("LDR", "Leg 1", "5V", "Connected to 5V power"),
    ("LDR", "Leg 2", "A3 + 10kΩ to GND", "Forms voltage divider"),
    ("LED 1", "Anode (+)", "D11 (via 220Ω)", "Node 1 indicator"),
    ("LED 2", "Anode (+)", "D10 (via 220Ω)", "Node 2 indicator"),
    ("LED 3", "Anode (+)", "D9 (via 220Ω)", "Node 3 indicator"),
    ("All LEDs", "Cathode (–)", "GND", "Common ground"),
    ("Servo 1", "Signal", "D5", "Node 1 servo control"),
    ("Servo 2", "Signal", "D6", "Node 2 servo control"),
    ("Servo 3", "Signal", "D3", "Node 3 servo control"),
    ("All Servos", "VCC / GND", "5V / GND", "Shared power and ground"),
    ("Buzzer", "Positive (+)", "D8", "Beep on quorum"),
    ("Buzzer", "Negative (–)", "GND", "Ground line"),
    ("Button", "Leg 1", "D2", "Reset input (INPUT_PULLUP)"),
    ("Button", "Leg 2", "GND", "Completes button circuit"),
    ("Arduino", "5V & GND", "Breadboard rails", "Distributes power")
]

for comp, leg, pin, desc in connections:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = leg
    row[2].text = pin
    row[3].text = desc

# Notes
doc.add_heading("Power and Connection Notes", level=2)
doc.add_paragraph(
    "All GND pins are connected together to ensure a common reference. "
    "The Arduino’s 5V pin powers potentiometers, LDR, and servos. "
    "If using multiple servos, use an external 5V 2A adapter with shared ground."
)

doc.save("Arduino_Wiring_Details_and_Pins.docx")
print("✅ Document created: Arduino_Wiring_Details_and_Pins.docx")
